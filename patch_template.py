import docx


def patch(file_path):
    doc = docx.Document(file_path)
    t = doc.tables[0]

    # 1. DC Title
    t.rows[0].cells[0].paragraphs[0].text = "{{ dc_title }}"

    # 2. DC No and Date
    t.rows[2].cells[2].paragraphs[0].text = "D.C No.          : {{ dc_no }}"
    t.rows[2].cells[2].paragraphs[1].text = "DATE              : {{ date }}"

    # 3. Deliver To
    t.rows[3].cells[2].paragraphs[0].text = "DELIVER TO:"
    if len(t.rows[3].cells[2].paragraphs) > 1:
        t.rows[3].cells[2].paragraphs[1].text = "{{ deliver_to }}"
    else:
        t.rows[3].cells[2].add_paragraph("{{ deliver_to }}")

    # 4. Items (Description & Qty)
    c_desc = t.rows[5].cells[1]
    c_qty = t.rows[5].cells[2]

    # Clear extra paragraphs in item cells
    for p in c_desc.paragraphs[1:]:
        p.text = ""
    for p in c_qty.paragraphs[1:]:
        p.text = ""

    # First cell opens the loop and renders description
    c_desc.paragraphs[0].text = "{% for item in items %}\n{{ item.description }}"

    # Second cell renders qty and closes the loop
    c_qty.paragraphs[0].text = "{{ item.qty }}\n{% endfor %}"

    doc.save(file_path)
    print("PATCHED:", file_path)


patch("/modules/assets/backend/templates/dc_template.docx")
